from docx import Document
import numpy

newfile = open('pdf.txt', 'w', encoding = 'utf8')
frase = "olá, não consigo acessar minha conta"
frase2 = frase.split(' ')
size = len(frase2)
prepos = ["olá", "ola", "oi", "bom", "dia", "boa", "tarde", "noite", "a", "o", "as", "os", "e", "ao", "aos", "à", "até", "não", "após", "ante", "com", "conforme", "contra", "de", "da", "do", "desde", "durante", "em", "entre", "mediante", "para", "perante", "por", "salvo", "sem", "sob", "sobre", "trás", "antes", "depois"]

matriz = numpy.zeros(shape=(255,1280))

paragraph = 0
paragraph_total = 0
j = 1
num_check = 0;

document = Document('..\pythonread\manual_rps_V19R01.docx')
full_size = len(document.paragraphs)
size_prepos = len(prepos)
print ("Os parágrafos são:\n")

for i in range (0,size):
    paragraph = 0;
    matriz [i][0] = i + 1
    j=1
    for para in document.paragraphs:
        for check_prepos in range (0,size_prepos):
            if frase2[i] == prepos[check_prepos]:
                num_check = 1
                break
        if num_check == 0:
            if frase2[i] in para.text:
                matriz [i][j] = paragraph
                resp = para.text + " - " + str(paragraph)
                print(resp.encode("utf-8"))
                newfile.write(para.text+"\n")
                j = j + 1
        else:
            num_check = 0
            break
        paragraph = paragraph + 1;
        

    print ("Fim para " + frase2[i] + "\n")

    if (paragraph_total <= j):
        paragraph_total = j

count = 0;
final_count = 0
final_paragraph = 0;
list = []
aux1 = 0;

for k in range (0, size):
    for m in range (1, paragraph_total):
        if matriz[k][m] != 0:
            for n in range (0, size):
                for o in range (1, paragraph_total):
                    if matriz[k][m] == matriz[n][o]:
                        count = count + 1
            if final_count < count:
                final_count = count
                final_paragraph = matriz[k][m]
                list.clear()
                list.append(final_paragraph)
            elif final_count == count:
                final_paragraph = matriz[k][m]
                aux0 = len(list)
                for a in range (0,aux0):
                    if list[a] == matriz[k][m]:
                        aux1 = 1
                if aux1 == 0:
                    list.append(final_paragraph)
                    aux1 = 0
            count = 0;

print (final_count)
print (list)

print (str(len(list))+"\n")
aux = len(list)

parag = 1
document = Document('..\pythonread\manual_rps_V19R01.docx')
for p in document.paragraphs:
    for q in range (0, aux):
        if parag == list[q] + 1:
            print(p.text)
    parag = parag + 1


# for q in range (0, aux):
#     pos = list[q]
#     print(str(pos)+ " - " + document.paragraphs[int(pos)].text + "\n")
